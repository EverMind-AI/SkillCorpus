"""GDPval deliverable grader — standalone LLM-judge rubric scorer.

Given a workspace directory containing an agent's deliverables and a GDPval
task (its official ``rubric_json`` + ``occupation``), an LLM judge inspects the
deliverables (Word/Excel/PDF/PowerPoint rendered to text or images, media read
via ffprobe) and scores each rubric criterion, returning a continuous reward in
[0, 1] equal to ``sum(score_i * met_i) / sum(positive scores)``.

The judge is any OpenAI-compatible chat model (default ``openai/gpt-4o`` via
OpenRouter; point ``base_url`` at ``https://api.openai.com/v1`` with model
``gpt-4o`` to use OpenAI directly). Configure via arguments or the environment
variables ``GDPVAL_API_KEY`` / ``OPENROUTER_API_KEY`` / ``OPENAI_API_KEY``,
``GDPVAL_JUDGE_MODEL`` (or legacy ``EVALUATION_MODEL``) and ``GDPVAL_BASE_URL``.

Lineage: the judge prompt + multimodal artifact rendering are adapted from
ClawWork's LLMEvaluator (https://github.com/HKUDS/ClawWork); the rubric-based
VERDICT scoring, artifact-reading fallbacks and degradation ladder are the
hardening added for GDPval grading.

System tools used when available (all degrade gracefully if absent):
  - libreoffice / soffice : recalc .xlsx formulas, render .pptx -> PDF
  - poppler (pdftoppm)    : render PDF pages to images (via pdf2image)
  - ffprobe (ffmpeg)      : read audio/video technical metadata
"""

import base64
import io
import json
import logging
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union

from openai import OpenAI

log = logging.getLogger("gdpval_grader")

# Per-occupation meta-prompts bundled with the package (fallback rubric source,
# used only when a task carries no rubric_json).
DEFAULT_META_PROMPTS_DIR = Path(__file__).resolve().parent / "meta_prompts"
DEFAULT_MODEL = "openai/gpt-4o"
DEFAULT_BASE_URL = "https://openrouter.ai/api/v1"

# ── File-reading helpers (from ClawWork livebench/tools/productivity/file_reading.py) ─

def _read_pdf_as_images(pdf_path: Path) -> Optional[List[bytes]]:
    """Convert PDF → list of PNG bytes (4 pages per combined 2×2 image)."""
    try:
        from PIL import Image
        from pdf2image import convert_from_path
    except ImportError:
        log.warning("pdf2image or Pillow not installed; falling back to text extraction for PDF")
        return None

    try:
        images = convert_from_path(str(pdf_path), dpi=100)
        if not images:
            return None

        combined_images = []
        for i in range(0, len(images), 4):
            batch = images[i:i + 4]
            resized = []
            for img in batch:
                if img.width > 600:
                    ratio = 600 / img.width
                    img = img.resize((600, int(img.height * ratio)), Image.Resampling.LANCZOS)
                resized.append(img)

            max_w = max(im.width for im in resized)
            max_h = max(im.height for im in resized)
            cols = 2
            rows = (len(resized) + 1) // 2
            combined = Image.new("RGB", (max_w * cols, max_h * rows), "white")
            for idx, img in enumerate(resized):
                combined.paste(img, ((idx % cols) * max_w, (idx // cols) * max_h))

            buf = io.BytesIO()
            combined.save(buf, format="PNG", optimize=True)
            combined_images.append(buf.getvalue())
        return combined_images
    except Exception as e:
        log.warning(f"PDF→image conversion failed: {e}")
        return None


def _read_pptx_as_images(pptx_path: Path) -> Optional[List[bytes]]:
    """Convert PPTX → list of PNG bytes (one per slide) via LibreOffice."""
    try:
        from PIL import Image
        from pdf2image import convert_from_path
    except ImportError:
        return None

    temp_dir = None
    try:
        temp_dir = tempfile.mkdtemp()
        # Private UserInstallation profile + HOME so concurrent graders don't
        # contend on the default ~/.config/libreoffice lockfile; cold-start can
        # exceed 30s under load, so allow 120s.
        _profile = os.path.join(temp_dir, "lo_profile")
        _env = dict(os.environ)
        _env["HOME"] = temp_dir
        result = subprocess.run(
            ["libreoffice", f"-env:UserInstallation=file://{_profile}",
             "--headless", "--convert-to", "pdf", "--outdir", temp_dir, str(pptx_path)],
            capture_output=True, timeout=120, text=True, env=_env,
        )
        if result.returncode != 0:
            return None

        pdf_name = pptx_path.stem + ".pdf"
        pdf_path = os.path.join(temp_dir, pdf_name)
        if not os.path.exists(pdf_path):
            return None

        images = convert_from_path(pdf_path, dpi=150)
        out = []
        for img in images:
            if img.width > 1200:
                ratio = 1200 / img.width
                img = img.resize((1200, int(img.height * ratio)), Image.Resampling.LANCZOS)
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            out.append(buf.getvalue())
        return out
    except Exception:
        return None
    finally:
        if temp_dir:
            shutil.rmtree(temp_dir, ignore_errors=True)


# ── Text-based fallback extractors ────────────────────────────────

def _salvage_text_artifact(path: str, size: int, kind: str) -> Dict:
    """Salvage an office file whose container is invalid (agents sometimes save
    plain text under a .docx/.xlsx name). Try reading it as UTF-8 text, then a
    same-stem .txt sibling; otherwise degrade to a present-but-unreadable note."""
    name = os.path.basename(path)
    try:
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        # NUL bytes decode fine as UTF-8 but mean binary garbage, not prose.
        if content.strip() and "\x00" not in content:
            return {"type": "text",
                    "content": f"[{kind} container invalid; raw text content of {name}]\n" + content}
    except UnicodeDecodeError:
        pass
    sibling = os.path.splitext(path)[0] + ".txt"
    if os.path.exists(sibling):
        try:
            with open(sibling, "r", encoding="utf-8") as f:
                content = f.read()
            if content.strip():
                return {"type": "text",
                        "content": f"[{kind} file {name} unreadable; grading same-stem "
                                   f"sibling {os.path.basename(sibling)}]\n" + content}
        except UnicodeDecodeError:
            pass
    return {"type": "note",
            "text": f"[{kind} file present but unreadable (invalid container): {name}, {size} bytes]",
            "size": size}


def _degrade_text_for_context(artifact_data: Dict, total_char_budget: int = 360000) -> Dict:
    """Bound the TOTAL text across all artifacts to fit the judge's context
    window. Per-file truncation alone is insufficient for many-file deliverables
    (e.g. 0353ee0c: dozens of files → 202k tokens even after per-file caps), so
    we enforce a global budget: split it evenly across text artifacts, then
    head/tail-truncate each. ~360k chars ≈ ~100k tokens, leaving room for the
    rubric/prompt under a 128k window."""
    out = dict(artifact_data)
    text_keys = [p for p, a in artifact_data.items() if a.get("type") == "text" and a.get("content")]
    if not text_keys:
        return out
    per_file = max(2000, total_char_budget // len(text_keys))
    for p in text_keys:
        c = artifact_data[p]["content"]
        if len(c) > per_file:
            half = per_file // 2
            out[p] = {"type": "text",
                      "content": c[:half]
                                 + f"\n...[{len(c) - per_file} chars truncated to fit judge context budget]...\n"
                                 + c[-half:]}
    return out


def _degrade_images_to_notes(artifact_data: Dict) -> Dict:
    """Replace image parts with notes when the judge API rejects the request
    (invalid_image_format zeroed 19 results)."""
    out = {}
    for p, art in artifact_data.items():
        if art.get("type") in ("image", "pptx_images", "pdf_images"):
            out[p] = {"type": "note",
                      "text": f"[visual artifact present but not accepted by the judge API: "
                              f"{os.path.basename(p)} ({art['type']})]",
                      "size": art.get("size", 0)}
        else:
            out[p] = art
    return out


def _cell_text_recursive(cell) -> str:
    """``Cell.text`` only joins the cell's own direct paragraphs — a cell whose
    content is itself a nested table (a common multi-column brochure/report
    layout: outer table for section slots, inner table per slot for the item
    list) reads back as blank/whitespace-only, silently dropping everything
    inside. Recurse into ``cell.tables`` so nested content is never invisible
    to the judge. Same bug class as the other "bad container -> fake 0"
    extraction gaps already patched in this module (see HANDOFF_GDPVAL.md).
    """
    parts = [p.text for p in cell.paragraphs if p.text.strip()]
    for nested in cell.tables:
        for row in nested.rows:
            parts.append(" | ".join(_cell_text_recursive(c) for c in row.cells))
    return "\n".join(parts).strip()


def _read_docx_content(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    doc = Document(path)
    parts = [f"[DOCX - {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables]"]
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for i, table in enumerate(doc.tables):
        parts.append(f"\n--- Table {i+1} ({len(table.rows)}×{len(table.columns)}) ---")
        for row in table.rows[:10]:
            parts.append(" | ".join(_cell_text_recursive(c) for c in row.cells))
        if len(table.rows) > 10:
            parts.append(f"... ({len(table.rows) - 10} more rows)")
    return "\n".join(parts)


def _recalc_xlsx_via_libreoffice(path: str) -> str:
    """Recalculate formulas via LibreOffice headless and return the path to a
    recalced copy. openpyxl-written files carry formulas but NO cached values,
    so ``load_workbook(data_only=True)`` reads every formula cell as None —
    blanking out the entire computed deliverable for the judge. LibreOffice
    opens + recomputes + re-saves so the cached values land on disk.

    Returns the original ``path`` unchanged on any failure (no libreoffice,
    timeout, convert error) so grading never gets *worse* than before.
    Parallel-safe: each call uses a private UserInstallation profile so
    concurrent graders don't contend on the LibreOffice lockfile.
    """
    import shutil as _shutil
    import tempfile as _tempfile
    import uuid as _uuid
    soffice = _shutil.which("libreoffice") or _shutil.which("soffice")
    if not soffice:
        return path
    try:
        outdir = _tempfile.mkdtemp(prefix="xlsx_recalc_")
        profile = os.path.join(outdir, "lo_profile")
        env = dict(os.environ)
        env["HOME"] = outdir  # avoid ~/.config/libreoffice lock contention
        subprocess.run(
            [soffice, f"-env:UserInstallation=file://{profile}",
             "--headless", "--calc", "--convert-to", "xlsx",
             "--outdir", outdir, path],
            check=True, capture_output=True, timeout=180, env=env,
        )
        out = os.path.join(outdir, os.path.splitext(os.path.basename(path))[0] + ".xlsx")
        return out if os.path.exists(out) else path
    except Exception:
        return path


def _read_xlsx_content(path: str) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ImportError("openpyxl not installed. Run: pip install openpyxl")
    # Recalc formulas first (openpyxl can't compute them; data_only reads None).
    path = _recalc_xlsx_via_libreoffice(path)
    wb = load_workbook(path, data_only=True)
    parts = [f"[Excel - {len(wb.sheetnames)} sheets: {', '.join(wb.sheetnames)}]"]
    for name in wb.sheetnames[:5]:
        ws = wb[name]
        parts.append(f"\n=== Sheet: {name} ({ws.max_row}×{ws.max_column}) ===")
        for ri, row in enumerate(ws.iter_rows(max_row=20, values_only=True), 1):
            row_text = " | ".join(str(c) if c is not None else "" for c in row)
            if row_text.strip():
                parts.append(f"Row {ri}: {row_text}")
        if (ws.max_row or 0) > 20:
            parts.append(f"... ({ws.max_row - 20} more rows)")
    if len(wb.sheetnames) > 5:
        parts.append(f"\n... ({len(wb.sheetnames) - 5} more sheets)")
    return "\n".join(parts)


_MEDIA_EXTENSIONS = {".mp4", ".mov", ".mkv", ".avi", ".webm", ".m4v",
                     ".wav", ".mp3", ".aac", ".flac", ".m4a", ".ogg", ".aiff", ".aif"}


def _read_media_metadata(path: str) -> str:
    """Extract technical media properties via ffprobe so the judge can verify
    codec/resolution/fps/sample-rate/bit-depth rubric criteria. Video/audio
    deliverables can't be inspected by an LLM directly; GDPval's media rubrics
    are written against exactly these ffprobe-checkable facts."""
    import subprocess
    try:
        out = subprocess.run(
            ["ffprobe", "-v", "error", "-print_format", "json",
             "-show_format", "-show_streams", path],
            capture_output=True, text=True, timeout=120)
        meta = json.loads(out.stdout or "{}")
    except Exception as e:
        return f"[media file {os.path.basename(path)}: ffprobe failed ({e})]"
    fmt = meta.get("format", {})
    lines = [f"[MEDIA technical metadata for {os.path.basename(path)} via ffprobe]",
             f"container/format: {fmt.get('format_name')}",
             f"duration_sec: {fmt.get('duration')}",
             f"size_bytes: {fmt.get('size')}", f"bit_rate: {fmt.get('bit_rate')}"]
    for s in meta.get("streams", []):
        ct = s.get("codec_type")
        if ct == "video":
            num, _, den = (s.get("r_frame_rate") or "0/1").partition("/")
            fps = (float(num) / float(den)) if den and float(den) else None
            lines.append(
                f"video stream: codec={s.get('codec_name')} ({s.get('codec_tag_string')}), "
                f"resolution={s.get('width')}x{s.get('height')}, "
                f"r_frame_rate={s.get('r_frame_rate')} (~{fps:.3f} fps)" if fps else
                f"video stream: codec={s.get('codec_name')}, resolution={s.get('width')}x{s.get('height')}, r_frame_rate={s.get('r_frame_rate')}")
            lines.append(f"  pix_fmt={s.get('pix_fmt')}, profile={s.get('profile')}, "
                         f"duration={s.get('duration')}, nb_frames={s.get('nb_frames')}")
        elif ct == "audio":
            lines.append(
                f"audio stream: codec={s.get('codec_name')}, sample_rate={s.get('sample_rate')} Hz, "
                f"channels={s.get('channels')}, sample_fmt={s.get('sample_fmt')}, "
                f"bits_per_raw_sample={s.get('bits_per_raw_sample')}, bits_per_sample={s.get('bits_per_sample')}")
    return "\n".join(lines)


# ── Core evaluator ────────────────────────────────────────────────

class LLMEvaluator:
    """Adapted from ClawWork — uses OpenRouter instead of OpenAI directly."""

    def __init__(self, meta_prompts_dir: str = "", model: str = "openai/gpt-4o",
                 api_key: str | None = None, base_url: str | None = None):
        # meta_prompts are only a FALLBACK for the rare task with no rubric_json;
        # default to the copy bundled with this package.
        self.meta_prompts_dir = Path(meta_prompts_dir) if meta_prompts_dir else DEFAULT_META_PROMPTS_DIR
        self.model = os.getenv("GDPVAL_JUDGE_MODEL") or os.getenv("EVALUATION_MODEL") or model
        api_key = (api_key or os.getenv("GDPVAL_API_KEY") or os.getenv("OPENROUTER_API_KEY")
                   or os.getenv("OPENAI_API_KEY") or "")
        if not api_key:
            raise ValueError(
                "No API key: pass api_key= or set GDPVAL_API_KEY / "
                "OPENROUTER_API_KEY / OPENAI_API_KEY")
        base_url = base_url or os.getenv("GDPVAL_BASE_URL") or "https://openrouter.ai/api/v1"

        self.client = OpenAI(base_url=base_url, api_key=api_key)
        self._cache: Dict[str, Dict] = {}
        log.info(f"LLMEvaluator: model={self.model} base_url={base_url}")

    # ── meta-prompt loading ──

    def _load_meta_prompt(self, occupation: str) -> Optional[Dict]:
        normalized = occupation.replace(" ", "_").replace(",", "")
        if normalized in self._cache:
            return self._cache[normalized]
        path = self.meta_prompts_dir / f"{normalized}.json"
        if not path.exists():
            log.warning(f"No meta-prompt for occupation '{occupation}' at {path}")
            return None
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
        self._cache[normalized] = data
        return data

    # ── artifact reading (multimodal) ──

    def _read_artifacts(self, paths: list[str]) -> Dict[str, Dict]:
        artifacts = {}
        for p in paths:
            size = os.path.getsize(p)
            ext = os.path.splitext(p)[1].lower()

            if size == 0:
                # Surface to the judge instead of zeroing the whole evaluation.
                artifacts[p] = {"type": "note",
                                "text": f"[file present but EMPTY (0 bytes): {os.path.basename(p)}]",
                                "size": 0}
                continue
            if ext in _MEDIA_EXTENSIONS:
                # Media: grade on ffprobe technical metadata (codec/resolution/fps/
                # sample-rate/bit-depth), which is what GDPval's media rubrics check.
                # Handled before the 2MB gate since media files are routinely larger.
                artifacts[p] = {"type": "note", "text": _read_media_metadata(p), "size": size}
                continue
            if size > 2 * 1024 * 1024:
                # Do NOT fail the whole evaluation over one oversized file (image-
                # heavy PPTX/PNG deliverables are common and legitimate). That
                # behaviour zeroed 227 results (13 tasks, ec-heavy). Instead,
                # surface the file to the judge as present-but-not-inlined.
                artifacts[p] = {"type": "note",
                                "text": f"[file present but too large to inline: "
                                        f"{os.path.basename(p)}, {size} bytes]",
                                "size": size}
                continue

            if ext in (".png", ".jpg", ".jpeg", ".gif", ".webp"):
                # Re-encode via PIL: validates the content AND normalizes the format.
                # Raw bytes with a lying extension (e.g. TIFF named .png) made the
                # judge API reject the whole request (invalid_image_format → 19
                # zeroed results). Unreadable images degrade to a note instead.
                try:
                    from PIL import Image
                    import io
                    with Image.open(p) as img:
                        img.load()
                        if img.mode not in ("RGB", "RGBA", "L"):
                            img = img.convert("RGB")
                        buf = io.BytesIO()
                        img.save(buf, format="PNG")
                    artifacts[p] = {"type": "image", "format": "png",
                                    "data": buf.getvalue(), "size": size}
                except Exception as e:
                    log.warning(f"Image unreadable for {p} ({e}); graded as note")
                    artifacts[p] = {"type": "note",
                                    "text": f"[image file present but unreadable/invalid: "
                                            f"{os.path.basename(p)}, {size} bytes]",
                                    "size": size}

            elif ext == ".pdf":
                images = _read_pdf_as_images(Path(p))
                if images:
                    artifacts[p] = {"type": "pdf_images", "images": images,
                                    "image_count": len(images), "size": size}
                else:
                    # Malformed/unconvertible PDF → fall back to text extraction so the
                    # task is still graded (don't crash the whole eval). "尽量都评".
                    txt = _read_pdf_text(p)
                    log.warning(f"PDF→image failed for {p}; graded via text fallback ({len(txt)} chars)")
                    artifacts[p] = {"type": "text", "content": txt}

            elif ext == ".pptx":
                images = _read_pptx_as_images(Path(p))
                if images:
                    artifacts[p] = {"type": "pptx_images", "images": images,
                                    "slide_count": len(images), "size": size}
                else:
                    # PPTX render failed → grade on a placeholder note rather than crash.
                    log.warning(f"PPTX render failed for {p}; graded as unrenderable note")
                    artifacts[p] = {"type": "text",
                                    "content": f"[PPTX file {os.path.basename(p)} ({size} bytes) "
                                               f"could not be rendered to images for inspection]"}

            elif ext == ".docx":
                # Invalid container (e.g. plain text saved as .docx) must not zero
                # the whole eval (PackageNotFoundError → 22 zeroed results). Try
                # plain-text salvage, then degrade to a note. "尽量都评".
                try:
                    artifacts[p] = {"type": "text", "content": _read_docx_content(p)}
                except Exception as e:
                    log.warning(f"DOCX read failed for {p} ({e}); trying text salvage")
                    artifacts[p] = _salvage_text_artifact(p, size, "DOCX")

            elif ext in (".xlsx", ".xls", ".xlsm"):
                try:
                    artifacts[p] = {"type": "text", "content": _read_xlsx_content(p)}
                except Exception as e:
                    log.warning(f"Excel read failed for {p} ({e}); trying text salvage")
                    artifacts[p] = _salvage_text_artifact(p, size, "Excel")

            else:
                try:
                    with open(p, "r", encoding="utf-8") as f:
                        artifacts[p] = {"type": "text", "content": f.read()}
                except UnicodeDecodeError:
                    artifacts[p] = {"type": "note",
                                    "text": f"[binary file present but not inlineable: "
                                            f"{os.path.basename(p)} ({ext}), {size} bytes]",
                                    "size": size}

        return artifacts

    # ── build multimodal content for API ──

    def _build_content(self, meta_prompt: Dict, task: Dict,
                       artifact_data: Dict, missing: list[str],
                       description: str, rubric: list | None = None) -> List[Dict]:
        # OFFICIAL alignment: when the task's published rubric_json is available,
        # grade against its per-task criteria (point-based). Otherwise fall back
        # to our per-occupation meta_prompt. The prompt wrapper + multimodal
        # artifact rendering below are unchanged (OpenAI's grader prompt isn't public).
        if rubric:
            evaluation_prompt = (
                "Grade the deliverable against the OFFICIAL rubric below. Each criterion has a "
                "SIGNED point value: POSITIVE = points earned if its stated condition is TRUE of the "
                "deliverable; NEGATIVE = a penalty that applies if its (undesirable) condition is TRUE.\n\n"
                "## Rubric criteria\n" + "\n".join(
                    f"[{i}] (score {c.get('score', 0):+d}) {c.get('criterion', '')}"
                    for i, c in enumerate(rubric)))
            category = "GDPval official per-task rubric"
        else:
            evaluation_prompt = meta_prompt.get("evaluation_prompt", "")
            category = meta_prompt.get('category', 'Unknown')

        text = f"""# TASK EVALUATION REQUEST

## Category: {category}

## Evaluation Guidelines:
{evaluation_prompt}

## Task Prompt (Original Assignment):
{task.get('prompt', 'N/A')}

## Task Metadata:
- Task ID: {task.get('task_id', 'N/A')}
- Sector: {task.get('sector', 'N/A')}
- Occupation: {task.get('occupation', 'N/A')}
- Reference Files: {', '.join(task.get('reference_files', [])) or 'None'}

## Agent's Description:
{description or 'No description provided'}

## Submitted Artifacts:

"""
        for path, art in artifact_data.items():
            name = os.path.basename(path)
            if art["type"] == "text":
                text += f"\n### File: {name}\n```\n{art['content']}\n```\n\n"
            elif art["type"] == "image":
                text += f"\n### Image: {name} ({art['format']}, {art['size']} bytes)\n[See image below]\n\n"
            elif art["type"] == "pptx_images":
                text += f"\n### PowerPoint: {name} ({art['slide_count']} slides)\n[See slide images below]\n\n"
            elif art["type"] == "pdf_images":
                approx = art["image_count"] * 4
                text += f"\n### PDF: {name} (~{approx} pages in {art['image_count']} combined images)\n[See PDF pages below]\n\n"
            elif art["type"] == "note":
                text += f"\n### File: {name}\n{art['text']}\n\n"

        if missing:
            text += "\n## Missing Artifacts:\n" + "".join(f"- {p}\n" for p in missing)

        if rubric:
            n = len(rubric)
            text += (
                "\n---\n\nThe criteria above are numbered [0] through "
                f"[{n - 1}]. For EACH criterion, decide whether its stated condition is TRUE of "
                "the deliverable (inspect the artifacts/images; do not credit unevidenced claims).\n\n"
                "Output your verdict for EVERY criterion as its own line, in EXACTLY this format "
                "(one line per criterion, all "
                f"{n} of them, indices 0 to {n - 1}):\n"
                "VERDICT[0]: 1\nVERDICT[1]: 0\n...\n"
                f"VERDICT[{n - 1}]: 1\n"
                "where the value is the single digit 1 if that criterion's condition is TRUE of the "
                "deliverable, else 0. You may add a brief justification after the digit on the same "
                f"line. Make sure you output all {n} VERDICT[i] lines.")
        else:
            text += """
---

Please evaluate this work according to the rubric above. Output your evaluation in this format:

**OVERALL SCORE:** [0-10]

**DIMENSION SCORES:**
[List dimension scores from rubric]

**KEY FINDINGS:**
[2-3 bullet points on what worked / didn't work]

**FEEDBACK:**
[1-2 paragraph explanation]

**TOP IMPROVEMENTS NEEDED:**
[Numbered list of 3 specific improvements]
"""
        content: List[Dict] = [{"type": "text", "text": text}]

        # Cap the number of inlined images. Many-image deliverables (a deck exported
        # page-by-page plus dozens of extracted crops) can exceed the grader API's
        # request-size / image-count limits; the API then returns a response with no
        # choices, and resp.choices[0] used to crash ('NoneType' object is not
        # subscriptable) and silently zero the whole eval. Send at most
        # _MAX_EVAL_IMAGES; surface the remainder as a text note.
        _MAX_EVAL_IMAGES = 16
        n_img = 0
        omitted = 0
        for path, art in artifact_data.items():
            if art["type"] == "image":
                if n_img >= _MAX_EVAL_IMAGES:
                    omitted += 1
                    continue
                b64 = base64.b64encode(art["data"]).decode()
                mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                        "gif": "image/gif", "webp": "image/webp"}.get(art["format"], "image/png")
                content.append({"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}", "detail": "high"}})
                n_img += 1

            elif art["type"] in ("pptx_images", "pdf_images"):
                for img_bytes in art["images"]:
                    if n_img >= _MAX_EVAL_IMAGES:
                        omitted += 1
                        continue
                    b64 = base64.b64encode(img_bytes).decode()
                    content.append({"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}", "detail": "high"}})
                    n_img += 1

        if omitted:
            content.append({"type": "text", "text":
                f"\n[Note: {omitted} additional image(s) present but not inlined "
                f"(grading image cap = {_MAX_EVAL_IMAGES}). Judge the criteria from the "
                f"{n_img} images shown plus the text artifacts above.]"})

        return content

    # ── score extraction ──

    @staticmethod
    def _extract_score(text: str) -> float:
        for pat in [r"OVERALL SCORE:\s*(\d+(?:\.\d+)?)",
                    r"Overall Score:\s*(\d+(?:\.\d+)?)",
                    r"Score:\s*(\d+(?:\.\d+)?)/10",
                    r"Final Score:\s*(\d+(?:\.\d+)?)"]:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                return max(0.0, min(10.0, float(m.group(1))))
        # fallback
        nums = re.findall(r"\b(\d+(?:\.\d+)?)\b", text[:200])
        for n in nums:
            v = float(n)
            if 0 <= v <= 10:
                return v
        log.warning("Could not extract score from evaluation, defaulting to 5.0")
        return 5.0

    @staticmethod
    def _parse_verdict(text: str, n: int) -> list | None:
        """Read the per-criterion `VERDICT[i]: 0/1` lines, mapping each value to its
        explicit index. Returns the 0/1 list iff ALL n indices 0..n-1 are present,
        else None.

        Indexed lines (vs one long comma array) are robust to the judge miscounting,
        reordering, or skipping: each value is pinned to its criterion index, so a
        missing one is detectable rather than silently shifting every later value.
        This replaces the old free-form met[] parsing (>=4 incompatible formats,
        bidirectional mis-scoring)."""
        met = LLMEvaluator._parse_verdict_partial(text, n)
        if any(v is None for v in met):
            return None
        return met

    @staticmethod
    def _parse_verdict_partial(text: str, n: int) -> list:
        """Like _parse_verdict but returns the partial list (None for any index the
        judge didn't provide) instead of None. Used to drive a targeted retry."""
        met = [None] * n
        for m in re.finditer(r'VERDICT\s*\[\s*(\d+)\s*\]\s*:?\s*([01])\b',
                             text or "", re.IGNORECASE):
            i = int(m.group(1))
            if 0 <= i < n:
                met[i] = int(m.group(2))   # last write wins if duplicated
        return met

    @staticmethod
    def _score_from_rubric(text: str, rubric: list) -> float | None:
        """Compute clamp( sum(score_i * met_i) / sum(positive scores), 0, 1 ) from
        the canonical FINAL_VERDICT line. Returns None if it can't be parsed
        (caller retries the judge, then flags — never silently scores 0)."""
        met = LLMEvaluator._parse_verdict(text, len(rubric))
        if met is None:
            return None
        pos = sum(c.get("score", 0) for c in rubric if c.get("score", 0) > 0)
        if not pos:
            return None
        total = sum(c.get("score", 0) * met[i] for i, c in enumerate(rubric))
        return max(0.0, min(1.0, total / pos))   # met-based → structurally <=1

    # ── main entry point ──

    def evaluate_artifact(self, task: Dict, artifact_paths: list[str],
                          description: str = "",
                          max_payment: float = 50.0) -> Tuple[float, str, float]:
        """Evaluate artifacts against occupation-specific rubric.

        Returns (normalized_score 0-1, feedback_text, payment).
        """
        occupation = task.get("occupation", "")
        if not occupation:
            return 0.0, "Error: no occupation", 0.0

        # OFFICIAL alignment: prefer the task's published rubric_json. meta_prompt
        # is only needed for the legacy fallback when no rubric is present.
        rubric = None
        try:
            rj = task.get("rubric_json")
            rubric = json.loads(rj) if isinstance(rj, str) else rj
            rubric = rubric if rubric else None
        except Exception:
            rubric = None

        meta_prompt = None
        if not rubric:
            meta_prompt = self._load_meta_prompt(occupation)
            if not meta_prompt:
                raise FileNotFoundError(f"No rubric_json and no meta-prompt for '{occupation}'")

        existing = [p for p in artifact_paths if os.path.exists(p)]
        missing = [p for p in artifact_paths if not os.path.exists(p)]
        if not existing:
            return 0.0, f"No artifacts found: {artifact_paths}", 0.0

        artifact_data = self._read_artifacts(existing)
        content = self._build_content(meta_prompt or {}, task, artifact_data, missing, description, rubric=rubric)

        sys_msg = ("You are an expert work evaluator. Follow the provided rubric precisely "
                   "and output a structured evaluation.")

        def _judge(extra_user: str = "") -> str:
            # content may be a plain str OR a multimodal list (text+image parts).
            if not extra_user:
                msg_content = content
            elif isinstance(content, list):
                msg_content = content + [{"type": "text", "text": extra_user}]
            else:
                msg_content = content + extra_user
            try:
                resp = self.client.chat.completions.create(
                    model=self.model,
                    temperature=0,
                    messages=[
                        {"role": "system", "content": sys_msg},
                        {"role": "user", "content": msg_content},
                    ],
                )
                # A rejected/oversized request (e.g. too many images) can return a
                # response with no choices; resp.choices[0] would then raise
                # 'NoneType' object is not subscriptable and get zeroed. Turn that
                # into a recognizable image error so the degradation ladder retries.
                if resp is None or not getattr(resp, "choices", None):
                    raise RuntimeError("grader returned no choices (request likely rejected: "
                                       "too many/unsupported images)")
                return resp.choices[0].message.content
            except Exception as e:
                raise RuntimeError(f"LLM evaluation failed: {e}") from e

        # Judge with degradation ladder instead of zeroing on API 400s:
        # 1st failure → truncate huge text artifacts and/or drop rejected images,
        # 2nd failure → both, last resort raises (caught upstream as error).
        try:
            eval_text = _judge()
        except RuntimeError as e:
            s = str(e)
            ctx_err = ("context_length_exceeded" in s or "maximum context length" in s
                       or "reduce the length" in s)
            img_err = ("invalid_image" in s or "unsupported image" in s.lower()
                       or "no choices" in s.lower() or "too many/unsupported images" in s.lower())
            if not (ctx_err or img_err):
                raise
            if ctx_err:
                artifact_data = _degrade_text_for_context(artifact_data)
            if img_err:
                artifact_data = _degrade_images_to_notes(artifact_data)
            log.warning("Judge API rejected request for task %s (%s); retrying degraded "
                        "(ctx_trunc=%s, img_drop=%s)",
                        task.get("task_id"), s[:120], ctx_err, img_err)
            content = self._build_content(meta_prompt or {}, task, artifact_data,
                                          missing, description, rubric=rubric)
            try:
                eval_text = _judge()
            except RuntimeError as e2:
                # Last rung: apply both degradations, then let failures surface.
                artifact_data = _degrade_images_to_notes(_degrade_text_for_context(artifact_data))
                content = self._build_content(meta_prompt or {}, task, artifact_data,
                                              missing, description, rubric=rubric)
                log.warning("Judge retry failed for task %s (%s); final fully-degraded attempt",
                            task.get("task_id"), str(e2)[:120])
                eval_text = _judge()

        if rubric:
            n = len(rubric)
            normalized = self._score_from_rubric(eval_text, rubric)
            # Never silently score 0: if any VERDICT[i] line is missing, RE-ASK the
            # judge once, naming exactly which indices are still needed.
            if normalized is None:
                partial = self._parse_verdict_partial(eval_text, n)
                missing = [i for i in range(n) if partial[i] is None]
                reminder = (f"\n\nIMPORTANT: your previous answer was missing verdicts for "
                            f"criteria {missing}. Output a line for EACH of those, exactly: "
                            f"VERDICT[i]: 1  (or 0). Provide all {len(missing)} missing lines now.")
                retry_text = _judge(reminder)
                # merge: prefer retry's values, fall back to first pass per index
                merged_partial = self._parse_verdict_partial(retry_text, n)
                merged = [merged_partial[i] if merged_partial[i] is not None else partial[i]
                          for i in range(n)]
                if all(v is not None for v in merged):
                    pos = sum(c.get("score", 0) for c in rubric if c.get("score", 0) > 0)
                    normalized = max(0.0, min(1.0,
                        sum(c.get("score", 0) * merged[i] for i, c in enumerate(rubric)) / pos)) if pos else None
                    eval_text = eval_text + "\n\n--- RETRY (missing verdicts) ---\n" + retry_text
                if normalized is None:
                    still = [i for i in range(n) if merged[i] is None]
                    log.warning("VERDICT[] incomplete after retry for task %s (missing %s); flagging",
                                task.get("task_id"), still)
                    normalized = 0.0
                    eval_text = f"[VERDICT incomplete after retry: missing {still}]\n" + (eval_text or "")
            # Degenerate-judgement tripwire: an all-zero verdict sheet with no
            # written justification is statistically a judge failure, not a real
            # all-fail (11 unflagged cases found 2026-06-05). Mark for analysis.
            if normalized == 0.0:
                ones = len(re.findall(r"VERDICT\[\d+\]:\s*1", eval_text or ""))
                prose = re.sub(r"VERDICT\[\d+\]:\s*\d", "", eval_text or "")
                if ones == 0 and len(prose.strip()) < 200:
                    eval_text = ("[JUDGING-DEGENERATE-SUSPECT: all-zero verdicts with no "
                                 "justification]\n" + (eval_text or ""))
        else:
            normalized = self._extract_score(eval_text) / 10.0   # legacy meta_prompt path
        payment = normalized * max_payment
        return normalized, eval_text, payment


def _read_pdf_text(path: str) -> str:
    """Fallback: extract text from PDF (when image conversion unavailable)."""
    try:
        import fitz
        doc = fitz.open(path)
        return "\n".join(page.get_text() for page in doc)
    except ImportError:
        pass
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    except ImportError:
        pass
    return f"[PDF file: {os.path.getsize(path)} bytes, no extraction library available]"


# ── Convenience wrapper for the domain adapter ─────────────────

_ARTIFACT_EXTENSIONS = {
    ".pdf", ".docx", ".xlsx", ".xls", ".xlsm", ".pptx",
    ".txt", ".csv", ".json", ".md", ".py", ".js", ".html", ".css",
    ".png", ".jpg", ".jpeg", ".gif", ".webp",
    # Media deliverables (graded on ffprobe technical metadata — see
    # _read_media_metadata; required for video/audio GDPval tasks).
    ".mp4", ".mov", ".mkv", ".avi", ".webm", ".m4v",
    ".wav", ".mp3", ".aac", ".flac", ".m4a", ".ogg", ".aiff", ".aif",
}

def grade_workspace(task: dict, workspace_dir, api_key: str | None = None,
                    meta_prompts_dir: str = "", model: str = DEFAULT_MODEL,
                    base_url: str | None = None,
                    evaluator: "LLMEvaluator | None" = None) -> dict:
    """Grade the deliverables in ``workspace_dir`` against a GDPval task's rubric.

    Args:
        task: dict with at least ``occupation`` and ``rubric_json`` (a list — or
            JSON string of a list — of ``{"criterion": str, "score": int}``).
            Optional: ``reference_files`` (basenames excluded from the graded
            deliverable set), ``prompt``, ``sector``, ``task_id``.
        workspace_dir: directory holding the agent's output files.
        evaluator: reuse an existing :class:`LLMEvaluator` (recommended when
            grading many tasks — it caches the OpenAI client and meta-prompts).
            If None, one is built from api_key/model/base_url/meta_prompts_dir.

    Returns:
        {"reward": float in [0,1], "evaluation_score": float, "score_raw": float,
         "feedback": str, "files_found": [str], ...}. ``reward`` is the continuous
        rubric score ``sum(score_i * met_i) / sum(positive scores)`` — NO cliff.
        On no gradable files: ``{"reward": 0.0, "error": "no_artifacts_found"}``.
    """
    workspace_dir = Path(workspace_dir)
    ev = evaluator or LLMEvaluator(meta_prompts_dir=meta_prompts_dir, model=model,
                                   api_key=api_key, base_url=base_url)

    # Discover artifacts in workspace (same logic as gdpval_bench)
    artifact_paths = []
    ref_basenames = {os.path.basename(r) for r in task.get("reference_files", [])}
    # Note: 'skills' excluded — paper-eval prompt-injection mode places retrieved
    # SKILL.md files under <workspace>/skills/<name>/. Those are inputs to the
    # agent (like reference_files), not agent deliverables. Including them in
    # the judge's artifact set would (a) contaminate the LLM judge with skill
    # text and (b) make the skill/noskill ablation asymmetric (skill condition
    # has skills/ dir, noskill condition doesn't).
    skip_dirs = {'node_modules', '__pycache__', '.git', 'venv', '.venv', 'skills'}
    skip_files = {'AGENTS.md', 'BOOTSTRAP.md', 'HEARTBEAT.md', 'IDENTITY.md',
                  'SOUL.md', 'TOOLS.md', 'USER.md', 'MEMORY.md',
                  'workspace-state.json'}
    if workspace_dir.exists():
        for f in sorted(workspace_dir.rglob("*")):
            if not f.is_file():
                continue
            if skip_dirs & set(f.relative_to(workspace_dir).parts):
                continue
            if f.suffix.lower() not in _ARTIFACT_EXTENSIONS:
                continue
            if f.name in ref_basenames or f.name in skip_files:
                continue
            if f.stat().st_size == 0:
                continue
            artifact_paths.append(str(f))

    if not artifact_paths:
        return {"reward": 0.0, "evaluation_score": 0.0,
                "error": "no_artifacts_found", "files_found": []}

    try:
        score, feedback, _payment = ev.evaluate_artifact(
            task=task,
            artifact_paths=artifact_paths,
            description=f"Work submission with {len(artifact_paths)} artifact(s)",
        )
    except Exception as e:
        log.error(f"  Evaluation failed: {e}")
        return {"reward": 0.0, "evaluation_score": 0.0, "error": str(e),
                "files_found": [os.path.basename(p) for p in artifact_paths]}

    # GDPval official rubric scoring is continuous — reward IS the rubric score.
    # (No payment cliff: zeroing scores in [0.3, 0.6) inflates variance and
    # creates spurious bimodality; it is a payment artifact, not part of grading.)
    return {
        "reward": round(score, 4),
        "evaluation_score": round(score, 4),
        "score_raw": round(score * 10, 2),
        "feedback": feedback,
        "files_found": [os.path.basename(p) for p in artifact_paths],
    }


# Backward-compatible name (the in-harness entry point was ``evaluate_rubric``).
def evaluate_rubric(task: dict, workspace_dir, api_key: str,
                    meta_prompts_dir: str = "", model: str = DEFAULT_MODEL) -> dict:
    return grade_workspace(task, workspace_dir, api_key=api_key,
                           meta_prompts_dir=meta_prompts_dir, model=model)
